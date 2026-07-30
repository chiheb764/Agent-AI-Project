from typing import TypedDict
import requests
from pypdf import PdfReader
from docx import Document
from langgraph.graph import StateGraph, END
import time


class AgentState(TypedDict):
    question: str
    reponse: str
    type_question: str


# def analyse_node(state):
#     print("Analyse de la question...")
#     return state


def analyse_node(state):
    question = state["question"]
    print("[LOG] Question reçue :", question)
    return state


def reponse_node(state):
    question = state["question"]
    state["reponse"] = f"Votre question est : {question}"
    return state


def decision_node(state):
    question = state["question"].lower()
    if "bonjour" in question:
        state["type_question"] = "salutation"
    elif "+" in question or "-" in question or "*" in question or "/" in question:

        state["type_question"] = "calcul"
    elif ".pdf" in question:
        state["type_question"] = "pdf"
    elif ".docx" in question:

        state["type_question"] = "docx"
    elif ".txt" in question:
        state["type_question"] = "txt"
    elif type(question) == str:
        state["type_question"] = "documentation"
    else:
        print("[LOG] Outil sélectionné :", state["type_question"])

    return state


def calculatrice_node(state):
    state["reponse"] = "Résultat du calcul"
    return state


# def documentation_node(state):
#     question = state["question"]
#     prompt = f"""
#     Réponds à cette question :
#     {question}
#     """
#     reponse = llm_local(prompt)
#     state["reponse"] = reponse
#     return state

# memoire = []
# memoire.append(f"Utilisateur : {question}")
# memoire.append(f"Assistant : {reponse}")
# historique = "\n".join(memoire)

memoire = []

historique = "\n".join(memoire)


def documentation_node(state):
    question = state["question"]
    prompt = f"""
    Historique :
    {historique}
    Question :
    {question}
    Réponse :
    """
    reponse = llm_local(prompt)
    state["reponse"] = reponse
    return state


def greeting_node(state):
    state["reponse"] = "Bonjour ! Comment puis-je vous aider ?"
    return state


# def pdf_reader_node(state):
#     contenu = pdf_reader("documents/formation.pdf")
#     question = state["question"]
#     prompt = f"""
#     Contexte :
#     {contenu}
#     Question :
#     {question}
#     Réponse :
#     """
#     state["reponse"] = llm_local(prompt)
#     return state


def pdf_reader_node(state):
    contenu = pdf_reader("documents/formation.pdf")
    question = state["question"]
    prompt = f"""
    Historique :
    {historique}
    Contexte :
    {contenu}
    Question :
    {question}
    Réponse :
    """
    state["reponse"] = llm_local(prompt)
    return state


# def docx_reader_node(state):
#     contenu = docx_reader("documents/procedure.docx")
#     state["reponse"] = contenu
#     return state


def docx_reader_node(state):
    contenu = docx_reader("documents/procedure.docx")
    question = state["question"]
    prompt = f"""
    Historique :
    {historique}
    Contexte :
    {contenu}
    Question :
    {question}
    Réponse :
    """
    state["reponse"] = llm_local(prompt)
    return state


def docx_reader_node(state):
    contenu = docx_reader("documents/procedure.docx")
    question = state["question"]

    prompt = f"""

    Contexte :
    {contenu}
    Question :
    {question}
    Réponse :
    """
    state["reponse"] = llm_local(prompt)
    return state


# def txt_reader(chemin_fichier):
#     with open(chemin_fichier, "r", encoding="utf-8") as fichier:
#         contenu = fichier.read()
#     return contenu


def txt_reader(chemin_fichier):
    try:
        with open(chemin_fichier, "r", encoding="utf-8") as fichier:
            return fichier.read()
    except:
        return "Fichier introuvable."


def route_question(state):
    return state["type_question"]


# def txt_reader_node(state):
#     contenu = txt_reader("documents/rh.txt")
#     question = state["question"]
#     prompt = f"""
#     Contexte :
#     {contenu}
#     Question :
#     {question}
#     Réponse :
#     """
#     state["reponse"] = llm_local(prompt)
#     return state


def txt_reader_node(state):
    contenu = txt_reader("documents/rh.txt")
    question = state["question"]
    prompt = f"""
    Historique :
    {historique}
    Contexte :
    {contenu}
    Question :
    {question}
    Réponse :
    """
    state["reponse"] = llm_local(prompt)
    return state


# def pdf_reader(chemin_fichier):
#     lecteur = PdfReader(chemin_fichier)
#     contenu = ""
#     for page in lecteur.pages:
#         contenu += page.extract_text()
#     return contenu


def pdf_reader(chemin_fichier):
    try:
        lecteur = PdfReader(chemin_fichier)
        contenu = ""
        for page in lecteur.pages:
            contenu += page.extract_text()
        return contenu
    except:
        return "Fichier introuvable."


# def docx_reader(chemin_fichier):
#     doc = Document(chemin_fichier)
#     contenu = ""
#     for paragraphe in doc.paragraphs:
#         contenu += paragraphe.text + "\n"
#     return contenu


def docx_reader(chemin_fichier):
    try:
        doc = Document(chemin_fichier)
        contenu = ""
        for paragraphe in doc.paragraphs:
            contenu += paragraphe.text + "\n"
        return contenu
    except:
        return "Fichier introuvable."


def llm_local(prompt):
    url = "http://localhost:11434/api/generate"
    data = {"model": "phi3", "prompt": prompt, "stream": False}
    response = requests.post(url, json=data)
    return response.json()["response"]


workflow = StateGraph(AgentState)

workflow.add_node("analyse", analyse_node)

workflow.add_node("reponse", reponse_node)

workflow.add_node("salutation", greeting_node)

workflow.add_node("decision", decision_node)
workflow.add_node("calculatrice", calculatrice_node)
workflow.add_node("documentation", documentation_node)


workflow.add_node("txt_reader", txt_reader_node)
workflow.add_node("pdf_reader", pdf_reader_node)
workflow.add_node("docx_reader", docx_reader_node)

workflow.add_conditional_edges(
    "decision",
    route_question,
    {
        "calcul": "calculatrice",
        "lecture": "txt_reader",
        "documentation": "documentation",
        "pdf": "pdf_reader",
        "docx": "docx_reader",
        "salutation": "salutation",
    },
)

workflow.set_entry_point("analyse")

workflow.add_edge("analyse", "decision")

workflow.add_edge("documentation", END)

workflow.add_edge("calculatrice", END)


workflow.add_edge("salutation", END)

workflow.add_edge("txt_reader", END)
workflow.add_edge("pdf_reader", END)
workflow.add_edge("docx_reader", END)

agent = workflow.compile()
# if question == "":
#     print("Veuillez saisir une question.")
resultat = agent.invoke({"question": "Lis formation.pdf"})
resultat2 = agent.invoke({"question": "Lis procedure.docx"})
# print(resultat)
# print(resultat2)
# # print(llm_local("Hello"))
# resultat = agent.invoke({"question": "What is an Agent IA ?"})
# print(resultat["reponse"])
contenu = txt_reader("documents/rh.txt")
prompt = f"""
Contexte :
{contenu}
Question :
Quels sont les congés ?
Réponse :
"""

resultat = agent.invoke({"question": "Lis formation.pdf"})
resultat = agent.invoke({"question": "Quels sujets sont étudiés ?"})
resultat = agent.invoke({"question": "Lis procedure.docx"})
resultat = agent.invoke({"question": "Que dit la procédure RH ?"})

# debut = time.time()
# resultat = agent.invoke(
# {"question": question}
# )
# fin = time.time()
# print(
# "Temps :",
# fin - debut,
# "secondes"
# )
if __name__ == "__main__":
    memoire = []
    questions = [
        "Quels sont les congés ?",
        "Lis formation.pdf",
        "50+20",
        "Lis procedure.docx",
    ]

    for question in questions:
        if question == "":
            print("Veuillez saisir une question.")
            continue
        historique = "\n".join(memoire)
        debut = time.time()
        resultat = agent.invoke({"question": question})
        fin = time.time()
        reponse = resultat["reponse"]
        memoire.append(f"Utilisateur : {question}")
        memoire.append(f"Assistant : {reponse}")
        print(reponse)
        print("[LOG] Réponse générée")
        print("Temps :", fin - debut, "secondes")
        print("-------------")